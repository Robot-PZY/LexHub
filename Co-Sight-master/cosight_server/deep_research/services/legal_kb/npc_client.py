"""全国人大法律法规库 (flk.npc.gov.cn) HTTP 客户端 — 参考 LaborAid / legal-mcp-server。"""

from __future__ import annotations

import re
import tempfile
import time
import uuid
from dataclasses import dataclass
from pathlib import Path

import requests

from app.common.logger_util import logger

_BASE = "https://flk.npc.gov.cn"
_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    ),
    "Referer": "https://flk.npc.gov.cn/fl.html",
    "Accept": "application/json, text/plain, */*",
    "Content-Type": "application/json",
}
_REQUEST_DELAY = 0.6


def _strip_html(text: str) -> str:
    return re.sub(r"<[^>]+>", "", text or "").strip()


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document

        doc = Document(str(path))
        lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return "\n".join(lines)
    except Exception as exc:
        logger.warning("DOCX extract failed for %s: %s", path, exc)
        return ""


@dataclass
class LawSearchHit:
    bbbs: str
    title: str
    publish_date: str | None
    effective_date: str | None
    authority: str | None
    law_type: str | None


@dataclass
class LawDocument:
    bbbs: str
    title: str
    publish_date: str | None
    effective_date: str | None
    authority: str | None
    law_type: str | None
    source_url: str
    text: str


class NpcLawClient:
    """访问 flk.npc.gov.cn 新版 /law-search/ API。"""

    def __init__(self, request_delay: float = _REQUEST_DELAY):
        self._delay = request_delay

    def search(self, keyword: str, *, page_size: int = 5, fuzzy: bool = False) -> list[LawSearchHit]:
        payload = {
            "searchContent": keyword,
            "searchRange": 1,
            "searchType": 2 if fuzzy else 1,
            "pageNum": 1,
            "pageSize": page_size,
        }
        resp = requests.post(
            f"{_BASE}/law-search/search/list",
            json=payload,
            headers=_HEADERS,
            timeout=30,
        )
        resp.raise_for_status()
        rows = resp.json().get("rows") or []
        hits: list[LawSearchHit] = []
        for row in rows:
            hits.append(
                LawSearchHit(
                    bbbs=row.get("bbbs", ""),
                    title=_strip_html(row.get("title", "")),
                    publish_date=row.get("gbrq"),
                    effective_date=row.get("sxrq"),
                    authority=row.get("zdjgName"),
                    law_type=row.get("flxz"),
                )
            )
        time.sleep(self._delay)
        return hits

    def fetch_document(self, bbbs: str, *, title_hint: str = "") -> LawDocument:
        detail_resp = requests.get(
            f"{_BASE}/law-search/search/flfgDetails",
            params={"bbbs": bbbs},
            headers=_HEADERS,
            timeout=60,
        )
        detail_resp.raise_for_status()
        detail = detail_resp.json().get("data") or {}

        download_resp = requests.get(
            f"{_BASE}/law-search/download/pc",
            params={"bbbs": bbbs, "format": "docx"},
            headers=_HEADERS,
            timeout=60,
        )
        download_resp.raise_for_status()
        download_data = download_resp.json().get("data") or {}
        file_url = download_data.get("url")
        if not file_url:
            raise ValueError(f"无法获取法规文件下载链接: {title_hint or bbbs}")

        file_resp = requests.get(file_url, headers=_HEADERS, timeout=120)
        file_resp.raise_for_status()

        title = detail.get("title") or title_hint or bbbs
        tmp_dir = Path(tempfile.gettempdir()) / "lexhub_npc_crawl"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        tmp_path = tmp_dir / f"{uuid.uuid4().hex[:12]}.docx"
        try:
            tmp_path.write_bytes(file_resp.content)
            text = _extract_docx_text(tmp_path)
        finally:
            tmp_path.unlink(missing_ok=True)

        if not text:
            raise ValueError(f"法规正文提取失败: {title}")

        time.sleep(self._delay)
        return LawDocument(
            bbbs=bbbs,
            title=str(title).strip(),
            publish_date=detail.get("gbrq"),
            effective_date=detail.get("sxrq"),
            authority=detail.get("zdjgName"),
            law_type=detail.get("flxz"),
            source_url=f"https://flk.npc.gov.cn/detail.html?bbbs={bbbs}",
            text=text.strip(),
        )

    def fetch_by_keyword(self, keyword: str) -> LawDocument:
        hits = self.search(keyword)
        if not hits:
            raise ValueError(f"未找到法规: {keyword}")
        chosen = hits[0]
        for hit in hits:
            if hit.title == keyword or keyword in hit.title:
                chosen = hit
                break
        return self.fetch_document(chosen.bbbs, title_hint=chosen.title)

    def search_law(self, query: str, *, limit: int = 5) -> list[dict]:
        """轻量检索：仅返回元数据，不下载全文（适合兜底引用）。"""
        hits = self.search(query, page_size=limit, fuzzy=True)
        return [
            {
                "source": "国家法律法规数据库",
                "id": hit.bbbs,
                "title": hit.title,
                "authority": hit.authority,
                "law_type": hit.law_type,
                "publish_date": hit.publish_date,
                "effective_date": hit.effective_date,
                "url": f"https://flk.npc.gov.cn/detail.html?bbbs={hit.bbbs}",
            }
            for hit in hits
        ]

    def get_law_article(self, bbbs: str, article_hint: str = "") -> dict | None:
        """获取法规全文并尝试定位条文（简化版 get_law_article）。"""
        doc = self.fetch_document(bbbs)
        if not article_hint:
            return {"title": doc.title, "content": doc.text[:4000], "source_url": doc.source_url}
        pattern = re.compile(re.escape(article_hint))
        match = pattern.search(doc.text)
        if not match:
            return {"title": doc.title, "content": doc.text[:4000], "source_url": doc.source_url}
        start = match.start()
        next_article = re.search(r"\n第[一二三四五六七八九十百千万零〇\d]+条", doc.text[start + 1 :])
        end = start + 1 + next_article.start() if next_article else min(len(doc.text), start + 3000)
        return {
            "title": f"{doc.title} {article_hint}",
            "content": doc.text[start:end].strip(),
            "source_url": doc.source_url,
        }
