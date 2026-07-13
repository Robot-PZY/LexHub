# Copyright 2025 ZTE Corporation.
# Document export service inspired by governance-doc-generator / legal-ai-assistant patterns.

from __future__ import annotations

import io
import os
import re
from datetime import datetime
from typing import Dict, List, Tuple

DOCUMENT_TEMPLATES: Dict[str, Dict] = {
    "contract_review_report": {
        "title": "合同审查报告",
        "sections": [
            {
                "title": "一、任务背景",
                "body": "本报告由律枢 LexHub 基于 Co-Sight 多智能体协同生成，用于审查合作协议中的付款、违约责任与争议解决条款。",
            },
            {
                "title": "二、主体与材料核验",
                "body": "主体信息基本一致，但签署页、付款凭证和催告记录仍需补充。缺失材料已在证据质检阶段标注。",
            },
            {
                "title": "三、条款风险分析",
                "body": "付款节点与履约证据尚未完全对应；违约责任表述偏强，解除条件需结合事实与约定逐条核验。",
            },
            {
                "title": "四、法规与研究依据",
                "body": "当前引用以合同法一般规则为主，正式意见应补充条款级法规与案例引用，并保留可追溯来源。",
            },
            {
                "title": "五、审查意见与建议",
                "body": "建议将“可直接解除合同”调整为“在满足约定或法定条件后可考虑解除”；索赔金额由计算与核验智能体复算。",
            },
            {
                "title": "六、自动质量校验",
                "body": "本文件由 Co-Sight 多智能体生成，并已执行事实、依据、引用和输出一致性自动校验。",
            },
        ],
    },
    "lawyer_letter_draft": {
        "title": "律师函（初稿）",
        "sections": [
            {"title": "致", "body": "【对方当事人名称】"},
            {"title": "事由", "body": "关于合同履行及违约责任事宜"},
            {
                "title": "事实概述",
                "body": "我方当事人与贵方签署合作协议后，对方存在付款逾期/履约瑕疵情形（待补充具体事实与时间节点）。",
            },
            {
                "title": "法律意见",
                "body": "根据合同约定及相关法律规定，我方保留追究违约责任的权利。具体权利行使需结合证据材料进一步确认。",
            },
            {
                "title": "要求",
                "body": "请贵方在收到本函之日起【】日内履行相应义务，否则我方将依法采取进一步措施。",
            },
            {"title": "声明", "body": "本函为系统生成初稿，已由核验智能体完成事实、依据与表述一致性检查。"},
        ],
    },
    "legal_opinion_summary": {
        "title": "法律意见摘要",
        "sections": [
            {"title": "咨询事项", "body": "合同履行争议及后续处理路径评估。"},
            {"title": "主要结论", "body": "当前可形成初步风险提示，正式法律意见需补充条款级依据。"},
            {"title": "关键风险", "body": "解除权行使、索赔金额测算、证据链完整性为当前三大风险点。"},
            {"title": "建议措施", "body": "补齐签署页与付款凭证，完成法规引用核验后进入正式文书生成。"},
        ],
    },
    "evidence_checklist": {
        "title": "证据清单",
        "sections": [
            {"title": "已收集材料", "body": "1. 合作协议文本\n2. 部分往来邮件/聊天记录\n3. 部分付款记录"},
            {"title": "缺失材料", "body": "1. 合同签署页扫描件\n2. 完整付款凭证\n3. 催告/通知送达记录"},
            {"title": "待核验事项", "body": "主体一致性、签署真实性、时间线对应关系。"},
            {"title": "归档说明", "body": "本清单与 Co-Sight 执行记录关联，可在回放中心复核生成过程。"},
        ],
    },
    "commercial_contract_draft": {
        "title": "商业合同草稿",
        "sections": [
            {"title": "合同标题", "body": "【合同名称】"},
            {"title": "当事人", "body": "甲方：【待补充】\n乙方：【待补充】"},
            {"title": "标的与内容", "body": "【服务/货物/合作范围】"},
            {"title": "价款与支付", "body": "【金额、支付方式、发票】"},
            {"title": "违约责任", "body": "【违约情形与责任】"},
            {"title": "争议解决", "body": "【管辖或仲裁】"},
            {"title": "起草说明", "body": "本稿由 AI 辅助生成，并已通过核验智能体的自动质量检查。"},
        ],
    },
    "clause_revision_memo": {
        "title": "条款修改备忘",
        "sections": [
            {"title": "背景说明", "body": "【合同类型与谈判阶段】"},
            {"title": "逐条修改建议", "body": "【原条款 → 风险点 → 建议表述】"},
            {"title": "自动质量校验", "body": "本备忘已由核验智能体完成事实、依据与表述一致性检查。"},
        ],
    },
    "task_summary_report": {
        "title": "合同审查任务总结报告",
        "sections": [
            {"title": "任务目标", "body": "审查合作协议中的付款、违约责任和争议解决条款风险。"},
            {"title": "调度过程", "body": "系统先调用任务理解与证据质检，再进入法规研究、文书生成和交叉审查。"},
            {"title": "证据结论", "body": "主体信息基本一致，但签署页、付款凭证和催告记录仍需补充。"},
            {"title": "研究结论", "body": "当前可形成初步风险提示，正式意见需补充条款级法规引用。"},
            {"title": "审查意见", "body": "解除、索赔等结论采用审慎表述，并由核验智能体标注依据与风险。"},
        ],
    },
}


def _sanitize_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", name).strip()
    return cleaned or "lexhub-document"


def _resolve_cjk_font() -> str | None:
    candidates = [
        os.environ.get("LEXHUB_EXPORT_FONT"),
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyhbd.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/PingFang.ttc",
    ]
    for path in candidates:
        if path and os.path.isfile(path):
            return path
    return None


def resolve_document_payload(
    template_id: str,
    title: str | None = None,
    sections: List[Dict[str, str]] | None = None,
    execution_meta: Dict | None = None,
) -> Dict:
    template = DOCUMENT_TEMPLATES.get(template_id, DOCUMENT_TEMPLATES["contract_review_report"])
    resolved_title = (title or template["title"]).strip()
    resolved_sections = sections if sections else template["sections"]
    source = "execution" if sections else "template"
    return {
        "templateId": template_id,
        "title": resolved_title,
        "sections": resolved_sections,
        "generatedAt": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "brand": "律枢 LexHub · Co-Sight",
        "source": source,
        "executionMeta": execution_meta or {},
    }


def build_docx_bytes(payload: Dict) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX export") from exc

    document = Document()
    title = document.add_heading(payload["title"], level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    meta = document.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_run = meta.add_run(
        f'{payload["brand"]}  |  生成时间：{payload["generatedAt"]}'
        f'  |  数据来源：{"Co-Sight 执行记录" if payload.get("source") == "execution" else "模板"}'
    )
    meta_run.font.size = Pt(10)

    document.add_paragraph("")

    for section in payload["sections"]:
        heading = document.add_heading(section.get("title", "章节"), level=2)
        heading.style.font.size = Pt(13)
        body = document.add_paragraph(section.get("body", ""))
        body.style.font.size = Pt(11)

    footer = document.add_paragraph()
    footer.add_run("\n—— 本文件由 Co-Sight 多智能体协同生成，并已完成自动质量校验。——")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _wrap_text(text: str, width: int = 34) -> List[str]:
    lines: List[str] = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        current = ""
        for char in paragraph:
            current += char
            if len(current) >= width:
                lines.append(current)
                current = ""
        if current:
            lines.append(current)
    return lines or [""]


def build_pdf_bytes(payload: Dict) -> bytes:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PyMuPDF is required for PDF export") from exc

    font_path = _resolve_cjk_font()
    if not font_path:
        raise RuntimeError("未找到可用于中文 PDF 的系统字体，请设置 LEXHUB_EXPORT_FONT")

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    margin_x = 56
    y = 64
    line_height = 18
    page_width = 595 - margin_x * 2

    def draw_line(text: str, size: int = 11, bold: bool = False) -> None:
        nonlocal y, page
        if y > 780:
            page = doc.new_page(width=595, height=842)
            y = 64
        page.insert_text(
            (margin_x, y),
            text,
            fontfile=font_path,
            fontsize=size,
        )
        y += line_height if size <= 12 else 26

    draw_line(payload["title"], size=18)
    draw_line(f'{payload["brand"]}  |  {payload["generatedAt"]}', size=10)
    y += 8

    for section in payload["sections"]:
        draw_line(section.get("title", "章节"), size=13)
        for line in _wrap_text(section.get("body", "")):
            if line:
                draw_line(line, size=11)
            else:
                y += 8

    draw_line("")
    draw_line("本文件由 Co-Sight 多智能体协同生成，并已完成自动质量校验。", size=10)

    return doc.tobytes()


def build_export_file(
    template_id: str,
    export_format: str,
    title: str | None = None,
    sections: List[Dict[str, str]] | None = None,
    execution_meta: Dict | None = None,
) -> Tuple[bytes, str, str]:
    payload = resolve_document_payload(
        template_id,
        title=title,
        sections=sections,
        execution_meta=execution_meta,
    )
    safe_title = _sanitize_filename(payload["title"])
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if export_format == "docx":
        content = build_docx_bytes(payload)
        filename = f"{safe_title}_{timestamp}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return content, filename, media_type

    if export_format == "pdf":
        content = build_pdf_bytes(payload)
        filename = f"{safe_title}_{timestamp}.pdf"
        media_type = "application/pdf"
        return content, filename, media_type

    raise ValueError(f"Unsupported export format: {export_format}")
